import os
import io
import csv
import requests
from typing import Annotated, Sequence
import logging
import sys
from datetime import datetime
from functools import lru_cache
import asyncio
import uvicorn
from fastapi import FastAPI, BackgroundTasks, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.runnables import RunnableConfig

# Core LangChain & LangGraph packages
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage

# REPLACED: ChatOllama with ChatGroq
from langchain_groq import ChatGroq

from langgraph.checkpoint.memory import MemorySaver
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from pydantic import BaseModel
from typing_extensions import TypedDict
from jina_embeddings import JinaEmbeddings

from dotenv import load_dotenv
load_dotenv()

# Setup logging to console and file
logging.basicConfig(
    level=logging.WARNING,
    format='[%(asctime)s] %(levelname)s: %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('chatbot_server.log')
    ]
)
logger = logging.getLogger(__name__)
logger.info("=" * 80)
logger.info("Starting eSpear Bot Initialization")
logger.info("=" * 80)

# ==========================================
# FASTAPI APP SETUP
# ==========================================
app = FastAPI(title="eSpear Bot", description="RAG-based chatbot backend")


# Request/Response models
class ChatRequest(BaseModel):
    message: str
    thread_id: str = "default_thread"


class ChatResponse(BaseModel):
    reply: str


# ==========================================
# CONNECT TO YOUR LOCAL CHROMA VECTOR DB
# ==========================================
persist_db_dir = "./chroma_db"
logger.info(f"Loading ChromaDB from: {persist_db_dir}")

# Setup the matching Nomic librarian model (smaller, faster embedding)
embeddings = JinaEmbeddings(
    api_key=os.environ["JINA_API_KEY"],
    model="jina-embeddings-v5-text-small"
)

logger.info("Embeddings model loaded: jina-embeddings-v5")

# Load existing Chroma database folder
vector_store = Chroma(
    persist_directory=persist_db_dir,
    embedding_function=embeddings,
    collection_metadata={"hnsw:space": "cosine"}  # Faster similarity search
)
logger.info("ChromaDB loaded successfully")

# Optimized retriever: k=2 + MMR (maximal marginal relevance) for diversity
retriever = vector_store.as_retriever(
    search_type="mmr",  # Better results with fewer docs
    search_kwargs={"k": 4, "lambda_mult": 0.5}
)
logger.info("Retriever initialized with MMR, k=2")

# ==========================================
# REPLACED: CHAT OLLAMA WITH CHAT GROQ
# ==========================================
# Replace "YOUR_GROQ_API_KEY" with your actual key or use os.environ.get("GROQ_API_KEY")
llm = ChatGroq(
    model="llama-3.3-70b-versatile",
    temperature=0.3,
    groq_api_key=os.environ.get("GROQ_API_KEY")
)
logger.info("LLM model loaded: llama-3.3-70b-versatile via Groq")
logger.info("=" * 80)

session_vectorstores: dict[str, Chroma] = {}
session_filenames: dict[str, list[str]] = {}
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)


def extract_text(filename: str, data: bytes) -> str:
    """Best-effort text extraction for the file types the uploader accepts."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    if ext == "docx":
        import docx
        document = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in document.paragraphs)

    if ext in ("txt", "md"):
        return data.decode("utf-8", errors="ignore")

    if ext == "csv":
        text = data.decode("utf-8", errors="ignore")
        rows = list(csv.reader(io.StringIO(text)))
        return "\n".join(", ".join(row) for row in rows)

    if ext in ("xlsx", "xls"):
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), data_only=True)
        lines = []
        for sheet in wb.worksheets:
            lines.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                lines.append(", ".join("" if c is None else str(c) for c in row))
        return "\n".join(lines)

    raise ValueError(f"Unsupported file type: .{ext}")


def get_session_context(thread_id: str, query: str, k: int = 3) -> str:
    vs = session_vectorstores.get(thread_id)
    if vs is None:
        return ""
    docs = vs.similarity_search(query, k=k)
    return "\n\n".join(doc.page_content for doc in docs)


# ==========================================
# RENDER LOG DIAGNOSIS
# ==========================================
# Reads recent logs from a service deployed on Render, picks out the error
# entries, and uses a self-authored troubleshooting guide (ingested
# separately via ingest_troubleshooting_kb.py into
# ./chroma_db_troubleshooting) as a knowledge base to suggest a fix for
# each one.
RENDER_API_KEY = os.environ.get("RENDER_API_KEY")
RENDER_OWNER_ID = os.environ.get("RENDER_OWNER_ID")
RENDER_DUMMY_SERVICE_ID = os.environ.get("RENDER_DUMMY_SERVICE_ID")

troubleshooting_vector_store = Chroma(
    persist_directory="./chroma_db_troubleshooting",
    embedding_function=embeddings,
    collection_name="troubleshooting_kb",
    collection_metadata={"hnsw:space": "cosine"},
)
troubleshooting_retriever = troubleshooting_vector_store.as_retriever(
    search_type="mmr",
    search_kwargs={"k": 4, "lambda_mult": 0.5},
)


def get_troubleshooting_context(query: str, k: int = 3, min_relevance: float = 0.6) -> str:
    """Only pull in troubleshooting KB context when it's actually relevant —
    otherwise every chat message (even unrelated ones) would get KB noise
    injected via plain top-k retrieval."""
    results = troubleshooting_vector_store.similarity_search_with_relevance_scores(query, k=k)
    relevant = [doc.page_content for doc, score in results if score >= min_relevance]
    return "\n\n".join(relevant)


def fetch_render_logs(service_id: str, limit: int = 100) -> list[dict]:
    if not (RENDER_API_KEY and RENDER_OWNER_ID):
        raise RuntimeError("RENDER_API_KEY / RENDER_OWNER_ID not configured")
    response = requests.get(
        "https://api.render.com/v1/logs",
        headers={"Authorization": f"Bearer {RENDER_API_KEY}"},
        params={"resource": service_id, "ownerId": RENDER_OWNER_ID, "limit": limit},
        timeout=30,
    )
    response.raise_for_status()
    return response.json().get("logs", [])


def _log_level(log: dict) -> str:
    return next((l["value"] for l in log.get("labels", []) if l["name"] == "level"), "")


def diagnose_render_errors(service_id: str, limit: int = 100, max_errors: int = 5) -> list[dict]:
    logs = fetch_render_logs(service_id, limit=limit)

    # Python logger.exception() writes the summary line at "error" level and
    # the traceback (which has the actual exception type/message we care
    # about) as separate "info"-level lines right after it. Stitch those back
    # together so the real error text is available for retrieval.
    error_logs = []
    i = 0
    while i < len(logs):
        if _log_level(logs[i]) == "error":
            parts = [logs[i]["message"]]
            j = i + 1
            while j < len(logs) and _log_level(logs[j]) != "error" and j - i <= 5:
                next_msg = logs[j]["message"]
                # Stop stitching once we hit the next real timestamped log line.
                if next_msg.startswith("[20") or _log_level(logs[j]) not in ("info", ""):
                    break
                parts.append(next_msg)
                j += 1
            log_copy = dict(logs[i])
            log_copy["message"] = "\n".join(parts)
            error_logs.append(log_copy)
            i = j
        else:
            # Skip "warning" entries — they're generic noise (slow query,
            # retry counters, queue depth) never meant to match the KB, and
            # they're frequent enough to crowd out the real exceptions in
            # the most-recent-N window.
            i += 1
    # Dedupe by exception type (the last line of the stitched traceback,
    # e.g. "DatabaseConnectionError: ...") so a run of repeats from the
    # small, randomly-sampled EXCEPTIONS pool doesn't crowd out variety —
    # keep the most recent occurrence of each distinct error type.
    seen_types = set()
    deduped = []
    for log in reversed(error_logs):
        exc_type = log["message"].splitlines()[-1].split(":")[0].strip()
        if exc_type in seen_types:
            continue
        seen_types.add(exc_type)
        deduped.append(log)
    error_logs = list(reversed(deduped))[-max_errors:]

    diagnoses = []
    for log in error_logs:
        message = log["message"]
        docs = troubleshooting_retriever.invoke(message)
        kb_context = "\n\n".join(f"[{d.metadata.get('source')}]\n{d.page_content}" for d in docs)

        prompt = f"""
        You are a debugging assistant. A service logged the following error:

        {message}

        Here is relevant guidance retrieved from the troubleshooting knowledge base:

        {kb_context}

        In 2-4 sentences, explain what's likely causing this error and how to
        fix it, based on the retrieved guidance. If the retrieved context
        doesn't cover it, say so plainly instead of guessing.
        """
        response = llm.invoke([HumanMessage(content=prompt)])
        diagnoses.append({
            "timestamp": log.get("timestamp"),
            "message": message,
            "diagnosis": response.content,
        })

    return diagnoses


# ==========================================
# DEFINE LANGGRAPH WORKFLOW
# ==========================================
# 1. Define the Chatbot State
class ChatState(TypedDict):
    messages: Annotated[Sequence[BaseMessage], add_messages]


# Query cache for common questions
@lru_cache(maxsize=128)
def get_cached_context(query: str):
    """Cache retrieval results for identical queries"""
    docs = retriever.invoke(query)
    return "\n\n".join([doc.page_content for doc in docs])


# 2. Define the Agent node execution - optimized
def chatbot_node(state: ChatState, config: RunnableConfig):
    user_message = state["messages"][-1].content
    thread_id = config.get("configurable", {}).get("thread_id", "default_thread")

    # Try cache first for faster response on repeated queries
    context = get_cached_context(user_message)

    # Blend in any document the user uploaded for this conversation
    session_context = get_session_context(thread_id, user_message)
    if session_context:
        context = f"{context}\n\n{session_context}" if context else session_context

    # Blend in the troubleshooting knowledge base — lets users paste an error
    # message directly into chat and get the same fix the log-diagnosis
    # feature would give.
    ts_context = get_troubleshooting_context(user_message)
    if ts_context:
        context = f"{context}\n\n{ts_context}" if context else ts_context

    # Shorter, more efficient system prompt
    system_prompt = f"""
    You are a helpful business assistant.

    Answer the user's question ONLY using the provided company documents.

    Instructions:
    - Use the context whenever possible.
    - If the context includes a troubleshooting entry (Error / Cause /
      Solution) that matches what the user described, give that solution
      directly.
    - If the answer is not present in the context, say:
      "I don't have enough information to answer that based on the available documents."
    - Do NOT make up facts.
    - Format answers clearly using bullet points or paragraphs when appropriate.

    Context:
    {context}
    """

    # Keep only last 2 messages for minimal context
    messages_window = state["messages"][-2:] if len(state["messages"]) > 2 else state["messages"]
    formatted_messages = [HumanMessage(content=system_prompt)] + list(messages_window)

    # Invoke the LLM
    response = llm.invoke(formatted_messages)
    return {"messages": [AIMessage(content=response.content)]}


# 3. Build and Compile the Workflow with Memory
workflow = StateGraph(ChatState)
workflow.add_node("chatbot", chatbot_node)
workflow.add_edge(START, "chatbot")
workflow.add_edge("chatbot", END)

# Memory saver allows the graph to remember multi-turn conversations
memory = MemorySaver()
app_graph = workflow.compile(checkpointer=memory)


# ==========================================
# API ENDPOINTS
# ==========================================
@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """Process chat message and return response from the RAG pipeline"""
    import time
    start_time = time.time()
    try:
        config = {"configurable": {"thread_id": request.thread_id}}

        # Execute graph with streaming
        events = app_graph.stream(
            {"messages": [HumanMessage(content=request.message)]},
            config
        )

        # Extract response efficiently
        response_text = ""
        for event in events:
            for value in event.values():
                if isinstance(value, dict) and 'messages' in value:
                    if value['messages']:
                        response_text = value['messages'][-1].content

        elapsed_time = time.time() - start_time
        logger.info(f"✓ {request.thread_id}: {elapsed_time:.2f}s")
        return ChatResponse(reply=response_text or "No response generated.")

    except Exception as e:
        elapsed_time = time.time() - start_time
        logger.error(f"✗ Error ({elapsed_time:.2f}s): {str(e)}")
        return ChatResponse(reply=f"❌ Error: {str(e)}")


@app.post("/upload")
async def upload_document(file: UploadFile = File(...), thread_id: str = Form("default_thread")):
    """Ingest an uploaded document into a per-conversation vector store so
    the chatbot can answer questions about it on the fly."""
    data = await file.read()
    try:
        text = extract_text(file.filename, data)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract any text from that file.")

    chunks = text_splitter.split_text(text)
    docs = [Document(page_content=chunk, metadata={"source": file.filename}) for chunk in chunks]

    vs = session_vectorstores.get(thread_id)
    if vs is None:
        vs = Chroma(embedding_function=embeddings, collection_name=f"session-{thread_id}")
        session_vectorstores[thread_id] = vs
    vs.add_documents(docs)
    session_filenames.setdefault(thread_id, []).append(file.filename)

    logger.info(f"Ingested '{file.filename}' for thread {thread_id}: {len(chunks)} chunks")
    return {"status": "ok", "filename": file.filename, "chunks": len(chunks)}


@app.delete("/upload/{thread_id}")
async def clear_uploaded_documents(thread_id: str):
    """Remove all uploaded documents for a conversation."""
    session_vectorstores.pop(thread_id, None)
    session_filenames.pop(thread_id, None)
    return {"status": "ok"}


@app.get("/diagnose-logs")
async def diagnose_logs(service_id: str = None, limit: int = 100, max_errors: int = 5):
    """Fetch recent logs from the dummy Render service, pull out the
    error entries, and diagnose each one against the troubleshooting
    knowledge base."""
    sid = service_id or RENDER_DUMMY_SERVICE_ID
    if not sid:
        raise HTTPException(status_code=400, detail="No service_id provided and RENDER_DUMMY_SERVICE_ID not set")
    try:
        diagnoses = diagnose_render_errors(sid, limit=limit, max_errors=max_errors)
    except requests.exceptions.HTTPError as e:
        raise HTTPException(status_code=502, detail=f"Render API error: {e}")
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))

    return {"service_id": sid, "count": len(diagnoses), "diagnoses": diagnoses}


@app.get("/health")
async def health():
    """Health check endpoint"""
    return {
        "status": "ok",
        "service": "eSpear Bot",
        "cache_size": get_cached_context.cache_info().currsize
    }


@app.post("/clear-cache")
async def clear_cache():
    """Clear the query cache to free memory"""
    get_cached_context.cache_clear()
    return {"status": "cache cleared", "message": "Query cache has been cleared"}


@app.get("/cache-stats")
async def cache_stats():
    """Get cache statistics"""
    stats = get_cached_context.cache_info()
    return {
        "hits": stats.hits,
        "misses": stats.misses,
        "currsize": stats.currsize,
        "maxsize": stats.maxsize,
        "hit_rate": stats.hits / (stats.hits + stats.misses) if (stats.hits + stats.misses) > 0 else 0
    }


@app.on_event("startup")
async def startup_event():
    logger.info("=" * 80)
    logger.info("🤖 eSpear Bot Started!")
    logger.info("FastAPI Server is running and ready to accept requests")
    logger.info("Documentation available at: http://localhost:8000/docs")
    logger.info("=" * 80)


if __name__ == "__main__":
    logger.info("=" * 80)
    logger.info("Starting Uvicorn server on 0.0.0.0:8080")
    logger.info("Press CTRL+C to stop")
    logger.info("=" * 80)
    port = int(os.environ.get("PORT", "8080"))
    print(f"Starting server on port {port}")
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info", timeout_keep_alive=120)
